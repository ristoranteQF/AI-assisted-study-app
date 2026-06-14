from __future__ import annotations

import io
import logging
from pathlib import Path

from docx import Document
from pypdf import PdfReader


logger = logging.getLogger("studybuddy.file")


SUPPORTED_EXTS = {".pdf", ".docx", ".txt", ".md"}


def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(data)
    if suffix == ".docx":
        return _extract_docx(data)
    if suffix in (".txt", ".md"):
        return data.decode("utf-8", errors="replace").strip()
    raise ValueError(
        f"Unsupported file type '{suffix}'. Supported: {', '.join(sorted(SUPPORTED_EXTS))}"
    )


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            logger.warning("Failed to extract a PDF page: %s", exc)
            text = ""
        pages.append(text)
    return "\n\n".join(p.strip() for p in pages if p.strip())


def _extract_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
